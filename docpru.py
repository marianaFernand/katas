from docx import Document
from docx.shared import Pt
from time import strftime, gmtime

document = Document()

document.add_heading('Telegrama', 0)
hoy = strftime("%d - %b - %y", gmtime())

fecha = document.add_paragraph(hoy)
fecha.aligment = 2

de = document.add_paragraph()
de.add_run('To: ').bold = True
de.add_run('Destinatario')

para = document.add_paragraph()
para.add_run('From: ').bold = True
para.add_run('Remitente')


document.add_heading('Mensaje', level=1)

table = document.add_table(rows=2, cols=1)
table.style = "LightShading"

style = table.style
font = style.font
font.name = "Courier"
font.size = Pt(12)
celda_morse = table.rows[0].cells[0]
celda_morse.text = "._.._"

celda_plano = table.rows[1].cells[0]
celda_plano.text = "hola"


document.save('demo.docx')